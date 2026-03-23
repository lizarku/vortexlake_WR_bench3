# -*- coding: utf-8 -*-
import matplotlib
matplotlib.use('Agg')
matplotlib.rcParams['font.family'] = 'DejaVu Sans'
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, tempfile

# ============================================================
# 1. 데이터 정의
# ============================================================

write_data = {
    'tar':     {'parse': 48376, 'serialize': 100244, 'compress': 10524, 'save': 183, 'total': 181028,
                'orig': 5690924543, 'ser_size': 7114035200, 'comp_size': 306762497,
                'ser_ratio': 125.0, 'zstd_saving': 95.7, 'total_saving': 94.6},
    'kryo':    {'parse': 46446, 'serialize': 12858, 'compress': 7800, 'save': 162, 'total': 77621,
                'orig': 5690924543, 'ser_size': 2841639885, 'comp_size': 278787788,
                'ser_ratio': 49.9, 'zstd_saving': 90.2, 'total_saving': 95.1},
    'parquet': {'parse': 46748, 'serialize': 71418, 'compress': 6147, 'save': 125, 'total': 124545,
                'orig': 5690924543, 'ser_size': 1728613619, 'comp_size': 216465236,
                'ser_ratio': 30.4, 'zstd_saving': 87.5, 'total_saving': 96.2},
    'avro':    {'parse': 46604, 'serialize': 28976, 'compress': 7278, 'save': 150, 'total': 85868,
                'orig': 5690924543, 'ser_size': 2972934910, 'comp_size': 252354858,
                'ser_ratio': 52.2, 'zstd_saving': 91.5, 'total_saving': 95.6},
}

read_data = {
    'tar':     {'step1': {'decomp': 89357, 'deser': 44546, 'total': 134719},
                'step2': {'decomp': 76888, 'deser': 44053, 'total': 121739}},
    'kryo':    {'step1': {'decomp': 38524, 'deser': 6854, 'total': 45715},
                'step2': {'decomp': 43527, 'deser': 6783, 'total': 50646}},
    'parquet': {
        'step1': {'decomp': 21716, 'deser': 27823, 'total': 50085},
        'step2': {'decomp': 20026, 'deser': 19453, 'total': 39669},
        'step3': {'decomp': 22214, 'deser': 9482, 'total': 31892},
        'step4': {'decomp': 27955, 'deser': 4609, 'total': 32780},
    },
    'avro':    {'step1': {'decomp': 39630, 'deser': 23578, 'total': 63554},
                'step2': {'decomp': 49081, 'deser': 17853, 'total': 67260}},
}

FORMATS = ['tar', 'kryo', 'parquet', 'avro']
FORMAT_LABELS = {'tar': 'Tar+JSON', 'kryo': 'Kryo', 'parquet': 'Parquet', 'avro': 'Avro'}

tmpdir = tempfile.mkdtemp()

# ============================================================
# 2. 차트 생성
# ============================================================

def style_chart(fig, ax):
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#CCCCCC')
    ax.spines['bottom'].set_color('#CCCCCC')
    ax.tick_params(colors='#555555', labelsize=9)
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, _: f'{x:,.0f}'))

def save_chart(fig, name):
    path = os.path.join(tmpdir, name)
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def chart_write_stacked():
    fig, ax = plt.subplots(figsize=(7, 3.8))
    x = np.arange(4)
    w = 0.55
    phases = ['parse', 'serialize', 'compress', 'save']
    phase_labels = ['Read + Parse', 'Serialize', 'Compress (zstd)', 'Save']
    phase_colors = ['#4472C4', '#ED7D31', '#A5A5A5', '#70AD47']
    bottom = np.zeros(4)
    for phase, label, color in zip(phases, phase_labels, phase_colors):
        vals = [write_data[f][phase] for f in FORMATS]
        ax.bar(x, vals, w, bottom=bottom, label=label, color=color, edgecolor='white', linewidth=0.5)
        bottom += vals
    ax.set_xticks(x)
    ax.set_xticklabels([FORMAT_LABELS[f] for f in FORMATS], fontsize=10)
    ax.set_ylabel('Time (ms)', fontsize=10, color='#555555')
    ax.set_title('Write Benchmark - Processing Time by Phase', fontsize=12, fontweight='bold', color='#333333', pad=12)
    ax.legend(loc='upper right', fontsize=8, framealpha=0.9)
    for i, f in enumerate(FORMATS):
        ax.text(i, bottom[i] + 2000, f'{write_data[f]["total"]:,.0f} ms', ha='center', va='bottom', fontsize=8, color='#333333')
    style_chart(fig, ax)
    return save_chart(fig, 'write_stacked.png')

def chart_write_size():
    fig, ax = plt.subplots(figsize=(7, 3.5))
    x = np.arange(4)
    w = 0.3
    ser = [write_data[f]['ser_size'] / 1024**3 for f in FORMATS]
    comp = [write_data[f]['comp_size'] / 1024**3 for f in FORMATS]
    ax.bar(x - w/2, ser, w, label='After Serialization', color='#4472C4', edgecolor='white')
    ax.bar(x + w/2, comp, w, label='After Compression (zstd)', color='#70AD47', edgecolor='white')
    ax.axhline(y=5690924543/1024**3, color='#C00000', linestyle='--', linewidth=1, label=f'Original ({5690924543/1024**3:.1f} GB)')
    ax.set_xticks(x)
    ax.set_xticklabels([FORMAT_LABELS[f] for f in FORMATS], fontsize=10)
    ax.set_ylabel('Size (GB)', fontsize=10, color='#555555')
    ax.set_title('Write Benchmark - Size Comparison', fontsize=12, fontweight='bold', color='#333333', pad=12)
    ax.legend(fontsize=8, framealpha=0.9)
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, _: f'{x:.1f}'))
    style_chart(fig, ax)
    return save_chart(fig, 'write_size.png')

def chart_read_step1():
    fig, ax = plt.subplots(figsize=(7, 3.5))
    x = np.arange(4)
    w = 0.3
    decomp = [read_data[f]['step1']['decomp'] for f in FORMATS]
    deser = [read_data[f]['step1']['deser'] for f in FORMATS]
    ax.bar(x - w/2, decomp, w, label='Decompress (zstd)', color='#A5A5A5', edgecolor='white')
    ax.bar(x + w/2, deser, w, label='Deserialize', color='#ED7D31', edgecolor='white')
    ax.set_xticks(x)
    ax.set_xticklabels([FORMAT_LABELS[f] for f in FORMATS], fontsize=10)
    ax.set_ylabel('Time (ms)', fontsize=10, color='#555555')
    ax.set_title('Read Benchmark - Step1 Full Column Read', fontsize=12, fontweight='bold', color='#333333', pad=12)
    ax.legend(fontsize=8, framealpha=0.9)
    for i, f in enumerate(FORMATS):
        t = read_data[f]['step1']['total']
        ax.text(i, max(decomp[i], deser[i]) + 2000, f'{t:,.0f} ms', ha='center', fontsize=8, color='#333333')
    style_chart(fig, ax)
    return save_chart(fig, 'read_step1.png')

def chart_parquet_pruning():
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    steps = ['56 col\n(Full)', '40 col', '20 col', '10 col']
    deser_vals = [read_data['parquet'][f'step{i}']['deser'] for i in range(1,5)]
    total_vals = [read_data['parquet'][f'step{i}']['total'] for i in range(1,5)]
    x = np.arange(4)
    w = 0.35
    ax.bar(x - w/2, total_vals, w, label='Total (decomp + deser)', color='#70AD47', edgecolor='white')
    ax.bar(x + w/2, deser_vals, w, label='Deserialize only', color='#ED7D31', edgecolor='white')
    ax.set_xticks(x)
    ax.set_xticklabels(steps, fontsize=10)
    ax.set_ylabel('Time (ms)', fontsize=10, color='#555555')
    ax.set_title('Parquet Column Pruning Effect', fontsize=12, fontweight='bold', color='#333333', pad=12)
    ax.legend(fontsize=8, framealpha=0.9)
    for i in range(4):
        pct = (1 - deser_vals[i] / deser_vals[0]) * 100
        if i > 0:
            ax.text(i + w/2, deser_vals[i] + 800, f'-{pct:.0f}%', ha='center', fontsize=8, color='#C00000', fontweight='bold')
    style_chart(fig, ax)
    return save_chart(fig, 'parquet_pruning.png')

def chart_optimized_read():
    # 6개 항목: tar, kryo, parquet(40col), parquet(20col), parquet(10col), avro
    labels = ['Tar+JSON\n(N/A)', 'Kryo\n(POJO Reuse)', 'Parquet\n(40col)', 'Parquet\n(20col)', 'Parquet\n(10col)', 'Avro\n(SpecificDatum)']
    baseline = [
        read_data['tar']['step1']['deser'],
        read_data['kryo']['step1']['deser'],
        read_data['parquet']['step1']['deser'],
        read_data['parquet']['step1']['deser'],
        read_data['parquet']['step1']['deser'],
        read_data['avro']['step1']['deser'],
    ]
    optimized = [
        read_data['tar']['step2']['deser'],
        read_data['kryo']['step2']['deser'],
        read_data['parquet']['step2']['deser'],
        read_data['parquet']['step3']['deser'],
        read_data['parquet']['step4']['deser'],
        read_data['avro']['step2']['deser'],
    ]
    fig, ax = plt.subplots(figsize=(8, 4))
    x = np.arange(len(labels))
    w = 0.3
    ax.bar(x - w/2, baseline, w, label='Baseline (56col)', color='#4472C4', edgecolor='white')
    ax.bar(x + w/2, optimized, w, label='Optimized', color='#70AD47', edgecolor='white')
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=7.5)
    ax.set_ylabel('Deserialize Time (ms)', fontsize=10, color='#555555')
    ax.set_title('Optimized Read - Deserialization Comparison', fontsize=12, fontweight='bold', color='#333333', pad=12)
    ax.legend(fontsize=8, framealpha=0.9)
    for i in range(len(labels)):
        pct = (1 - optimized[i] / baseline[i]) * 100
        if abs(pct) > 1:
            ax.text(i + w/2, optimized[i] + 600, f'-{pct:.0f}%', ha='center', fontsize=7.5, color='#C00000', fontweight='bold')
    style_chart(fig, ax)
    return save_chart(fig, 'optimized_read.png')

def chart_total_rw():
    fig, ax = plt.subplots(figsize=(7, 3.5))
    x = np.arange(4)
    w = 0.3
    wr = [write_data[f]['total'] for f in FORMATS]
    rd = [read_data[f]['step1']['total'] for f in FORMATS]
    ax.bar(x - w/2, wr, w, label='Write', color='#4472C4', edgecolor='white')
    ax.bar(x + w/2, rd, w, label='Read', color='#ED7D31', edgecolor='white')
    ax.set_xticks(x)
    ax.set_xticklabels([FORMAT_LABELS[f] for f in FORMATS], fontsize=10)
    ax.set_ylabel('Time (ms)', fontsize=10, color='#555555')
    ax.set_title('Write vs Read - Total Time Comparison', fontsize=12, fontweight='bold', color='#333333', pad=12)
    ax.legend(fontsize=8, framealpha=0.9)
    style_chart(fig, ax)
    return save_chart(fig, 'total_rw.png')

chart_paths = {
    'write_stacked': chart_write_stacked(),
    'write_size': chart_write_size(),
    'read_step1': chart_read_step1(),
    'optimized_read': chart_optimized_read(),
    'parquet_pruning': chart_parquet_pruning(),
    'total_rw': chart_total_rw(),
}

# ============================================================
# 3. DOCX 생성
# ============================================================
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    hs.font.name = 'Calibri'

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    run.italic = True

# ============================================================
# 표지
# ============================================================
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('VortexLake')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u'\uc9c1\ub82c\ud654 \ud3ec\ub9f7 \uc131\ub2a5 \ubca4\uce58\ub9c8\ud06c \ubcf4\uace0\uc11c')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tar+JSON / Kryo / Parquet / Avro')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026-03-20  |  VortexLake Infrastructure Team')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# 목차
# ============================================================
doc.add_heading(u'\ubaa9\ucc28', level=1)
toc_items = [
    u'1. \uac1c\uc694',
    u'2. \ud14c\uc2a4\ud2b8 \ud658\uacbd',
    u'3. \uc4f0\uae30 \ubca4\uce58\ub9c8\ud06c',
    u'   3.1 \ucc98\ub9ac \uc2dc\uac04',
    u'   3.2 \uc6a9\ub7c9 \ud6a8\uc728',
    u'4. \uc77d\uae30 \ubca4\uce58\ub9c8\ud06c',
    u'   4.1 \uc804\uccb4 \ucee8\ub7fc \uc77d\uae30 (Step 1)',
    u'   4.2 \ubaa8\ub4c8\ubcc4 \ucd5c\uc801\ud654 \uc77d\uae30 \ube44\uad50',
    u'   4.3 Parquet \ucee8\ub7fc \ud504\ub8e8\ub2dd \uc2ec\ud654',
    u'5. \uc4f0\uae30/\uc77d\uae30 \uc885\ud569 \ube44\uad50',
    u'6. \ubc1c\uc0dd \uc774\uc288 \ubc0f \uc870\uce58 \uc0ac\ud56d',
    u'7. \uacb0\ub860',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. 개요
# ============================================================
doc.add_heading(u'1. \uac1c\uc694', level=1)

doc.add_heading(u'1.1 \ubca4\uce58\ub9c8\ud06c \ubaa9\uc801', level=2)
doc.add_paragraph(
    u'VortexLake \ub85c\uadf8 \uc800\uc7a5 \ud30c\uc774\ud504\ub77c\uc778\uc5d0\uc11c \uc0ac\uc6a9 \uac00\ub2a5\ud55c '
    u'4\uac00\uc9c0 \uc9c1\ub82c\ud654 \ud3ec\ub9f7(Tar+JSON, Kryo, Parquet, Avro)\uc758 \uc131\ub2a5\uc744 '
    u'\uc2e4 \uc6b4\uc601 \ud658\uacbd\uacfc \ucd5c\ub300\ud55c \ub3d9\uc77c\ud55c \uc870\uac74\uc5d0\uc11c \uce21\uc815\ud558\uc5ec, '
    u'\ud3ec\ub9f7 \uc120\uc815\uc5d0 \ud544\uc694\ud55c \uc815\ub7c9\uc801 \uadfc\uac70\ub97c \uc81c\uacf5\ud558\ub294 \uac83\uc774 \ubaa9\uc801\uc774\ub2e4.'
)

doc.add_heading(u'1.2 \ud14c\uc2a4\ud2b8 \ubc94\uc704', level=2)
doc.add_paragraph(
    u'\uc4f0\uae30 \uacbd\ub85c(Read + Parse \u2192 Serialize \u2192 Compress \u2192 Save)\uc640 '
    u'\uc77d\uae30 \uacbd\ub85c(Decompress \u2192 Deserialize)\ub97c \ubaa8\ub450 \ud3ec\ud568\ud558\uba70, '
    u'\ubaa8\ub4e0 \ud3ec\ub9f7\uc5d0 \ub3d9\uc77c\ud55c \ub370\uc774\ud130\uc14b, JVM \uc124\uc815, \uc555\ucd95 \uc54c\uace0\ub9ac\uc998, '
    u'I/O \ud328\ud134\uc744 \uc801\uc6a9\ud558\uc5ec \uacf5\uc815\ud55c \ube44\uad50\ub97c \ubcf4\uc7a5\ud558\uc600\ub2e4. '
    u'\uac01 \ud14c\uc2a4\ud2b8\ub294 JIT \uc6c0\uc5c5 \ud6c4 5~11\ud68c \ubc18\ubcf5 \uce21\uc815\ud558\uc5ec \uc548\uc815\uc801\uc778 \uacb0\uacfc\ub97c \ub3c4\ucd9c\ud558\uc600\ub2e4.'
)

# ============================================================
# 2. 테스트 환경
# ============================================================
doc.add_heading(u'2. \ud14c\uc2a4\ud2b8 \ud658\uacbd', level=1)

add_table(doc,
    [u'\ud56d\ubaa9', u'\uc0ac\uc591'],
    [
        ['JDK', 'Azul Zulu 21.0.2 (--enable-preview)'],
        [u'\ud799 \uba54\ubaa8\ub9ac', u'4,096 MB (-Xms4g -Xmx4g, \uace0\uc815)'],
        ['OS', 'Oracle Linux 8 (Kernel 5.15)'],
        [u'\uc555\ucd95', 'zstd Level 1 (zstd-jni 1.5.6-8)'],
        [u'\ub370\uc774\ud130\uc14b', u'1,086\uac1c .done \ub85c\uadf8 \ud30c\uc77c, \ucd1d 5.3 GB'],
        [u'\ub808\ucf54\ub4dc \uc218', u'4,741,381\uac74 (RegularLog, 56\uac1c \ud544\ub4dc)'],
        [u'\uc4f0\uae30 \ubc18\ubcf5', u'\ud3ec\ub9f7\ub2f9 11\ud68c'],
        [u'\uc77d\uae30 \ubc18\ubcf5', u'\ud3ec\ub9f7/\ub2e8\uacc4\ub2f9 5\ud68c'],
        [u'\uc6c0\uc5c5', u'30\uac1c \ud30c\uc77c (~138,000\uac74) / \uc2e4\ud589\ub2f9'],
    ]
)

doc.add_paragraph()
doc.add_heading(u'\ub370\uc774\ud130\uc14b \uc120\uc815 \uc0ac\uc720', level=3)
doc.add_paragraph(
    u'\ud569\uc131 \ub370\uc774\ud130\uac00 \uc544\ub2cc \uc2e4\uc81c \uc6b4\uc601 \ud658\uacbd\uc758 \ubcf4\uc548 \ub85c\uadf8 \ud30c\uc77c(.done \ud3ec\ub9f7)\uc744 \uc0ac\uc6a9\ud558\uc600\ub2e4. '
    u'\uc2e4\uc81c \ub370\uc774\ud130\ub97c \uc0ac\uc6a9\ud574\uc57c \ud544\ub4dc \ubd84\ud3ec, null \ube44\uc728, \ubb38\uc790\uc5f4 \uae38\uc774 \ub4f1\uc774 '
    u'\uc6b4\uc601 \ud658\uacbd\uc744 \ubc18\uc601\ud558\uba70, \ucd1d \uc6a9\ub7c9 5.3 GB(\uc57d 474\ub9cc \uac74)\ub294 '
    u'\uc2e4\uc81c \ubc30\uce58 \ucc98\ub9ac \ucc3d(batch window) \uaddc\ubaa8\uc5d0 \ud574\ub2f9\ud55c\ub2e4.'
)

doc.add_heading(u'JVM \uc6c0\uc5c5 \uc804\ub7b5', level=3)
doc.add_paragraph(
    u'\uac01 \uc2e4\ud589\ub9c8\ub2e4 \uce21\uc815 \uc804 30\uac1c \ud30c\uc77c(\uc57d 138,000\uac74)\uc744 \uc6c0\uc5c5\uc73c\ub85c \ucc98\ub9ac\ud55c\ub2e4. '
    u'JVM JIT \ucef4\ud30c\uc77c \uc218\uba85 \uc8fc\uae30\ub97c \uace0\ub824\ud55c \uac83\uc73c\ub85c, '
    u'C1 \ucef4\ud30c\uc77c\uc740 \uc57d 1,500\ud68c \ud638\ucd9c \uc2dc, C2\ub294 \uc57d 10,000\ud68c \uc2dc \ud2b8\ub9ac\uac70\ub418\uba70, '
    u'\uc774\ud6c4 deoptimization/\uc7ac\ucef4\ud30c\uc77c \uc8fc\uae30\uae4c\uc9c0 \uc644\ub8cc\ub418\uc5b4\uc57c \uc548\uc815\uc801\uc778 \uce21\uc815\uc774 \uac00\ub2a5\ud558\ub2e4. '
    u'30\uac1c \ud30c\uc77c \uc2dc\uc810\uc5d0\uc11c JIT\uac00 \uc644\uc804\ud788 \uc548\uc815\ud654\ub418\uba70, '
    u'\uc774\ub294 \uc7a5\uc2dc\uac04 \uc6b4\uc601\ub418\ub294 \ud504\ub85c\ub355\uc158 \ub370\ubaac\uc758 \uc0c1\ud0dc\uc640 \uc77c\uce58\ud55c\ub2e4. '
    u'4\uac1c \ud3ec\ub9f7 \ubaa8\ub450 \ub3d9\uc77c\ud55c \uc6c0\uc5c5 \ud69f\uc218\ub97c \uc801\uc6a9\ud558\uc5ec \uacf5\uc815\uc131\uc744 \uc720\uc9c0\ud558\uc600\ub2e4.'
)

doc.add_page_break()

# ============================================================
# 3. 쓰기 벤치마크
# ============================================================
doc.add_heading(u'3. \uc4f0\uae30 \ubca4\uce58\ub9c8\ud06c', level=1)
doc.add_paragraph(
    u'\uc4f0\uae30 \ud30c\uc774\ud504\ub77c\uc778\uc740 \ub85c\uadf8 \ud30c\uc77c\uc744 4\uac1c \ub2e8\uacc4\ub85c \uc21c\ucc28 \ucc98\ub9ac\ud55c\ub2e4: '
    u'\uc77d\uae30+\ud30c\uc2f1, \uc9c1\ub82c\ud654, \uc555\ucd95(zstd L1), \uc800\uc7a5. '
    u'\uac01 \ub2e8\uacc4\ub294 System.nanoTime() \uae30\ubc18\uc73c\ub85c \ub3c5\ub9bd\uc801\uc73c\ub85c \uc2dc\uac04\uc744 \uce21\uc815\ud558\uc600\ub2e4.'
)

# 3.1
doc.add_heading(u'3.1 \ucc98\ub9ac \uc2dc\uac04', level=2)
doc.add_picture(chart_paths['write_stacked'], width=Inches(5.8))
add_caption(doc, u'\uadf8\ub9bc 1. \uc4f0\uae30 \ud30c\uc774\ud504\ub77c\uc778 \ub2e8\uacc4\ubcc4 \uc18c\uc694\uc2dc\uac04 (11\ud68c \ud3c9\uade0)')

doc.add_paragraph()
add_table(doc,
    [u'\ud3ec\ub9f7', u'\uc77d\uae30+\ud30c\uc2f1', u'\uc9c1\ub82c\ud654', u'\uc555\ucd95', u'\uc800\uc7a5', u'\uc804\uccb4'],
    [
        [FORMAT_LABELS[f],
         f'{write_data[f]["parse"]:,} ms',
         f'{write_data[f]["serialize"]:,} ms',
         f'{write_data[f]["compress"]:,} ms',
         f'{write_data[f]["save"]:,} ms',
         f'{write_data[f]["total"]:,} ms']
        for f in FORMATS
    ]
)
add_caption(doc, u'\ud45c 1. \uc4f0\uae30 \ubca4\uce58\ub9c8\ud06c \ud3c9\uade0 \uc2dc\uac04 (ms)')

doc.add_paragraph()
doc.add_paragraph(
    u'\uc77d\uae30+\ud30c\uc2f1 \uc2dc\uac04\uc740 \ubaa8\ub4e0 \ud3ec\ub9f7\uc5d0\uc11c \uac70\uc758 \ub3d9\uc77c\ud558\ub2e4(\uc57d 47\ucd08). '
    u'\uc774\ub294 \uacf5\ud1b5 LogReader\ub97c \uc0ac\uc6a9\ud558\uae30 \ub54c\ubb38\uc774\uba70, \ud3ec\ub9f7 \uac04 \ucc28\uc774\ub294 \uc9c1\ub82c\ud654 \ub2e8\uacc4\uc5d0\uc11c \ubc1c\uc0dd\ud55c\ub2e4.'
)

bullets = [
    u'Kryo: \uac00\uc7a5 \ube60\ub978 \uc9c1\ub82c\ud654(12.9\ucd08). \uc2a4\ud0a4\ub9c8 \uc624\ubc84\ud5e4\ub4dc \uc5c6\ub294 \ucf64\ud329\ud2b8 \ubc14\uc774\ub108\ub9ac \uc778\ucf54\ub529.',
    u'Avro: 29.0\ucd08. \uc2a4\ud0a4\ub9c8 \uae30\ubc18 \ubc14\uc774\ub108\ub9ac \uc778\ucf54\ub529(ReflectDatumWriter).',
    u'Parquet: 71.4\ucd08. \ucee8\ub7fc\ubcc4 \ub515\uc154\ub108\ub9ac/RLE \uc778\ucf54\ub529\uc73c\ub85c \uc778\ud574 \ub290\ub9bc.',
    u'Tar+JSON: 100.2\ucd08. JSON \ud14d\uc2a4\ud2b8 \uc9c1\ub82c\ud654\uc758 \ub192\uc740 \ucd9c\ub825\ub7c9\uc73c\ub85c \uac00\uc7a5 \ub290\ub9bc.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# 3.2
doc.add_heading(u'3.2 \uc6a9\ub7c9 \ud6a8\uc728', level=2)
doc.add_picture(chart_paths['write_size'], width=Inches(5.8))
add_caption(doc, u'\uadf8\ub9bc 2. \uc9c1\ub82c\ud654/\uc555\ucd95 \ud6c4 \uc6a9\ub7c9 \ube44\uad50')

doc.add_paragraph()
add_table(doc,
    [u'\ud3ec\ub9f7', u'\uc6d0\ubcf8', u'\uc9c1\ub82c\ud654 \ud6c4', u'\uc9c1\ub82c\ud654 \ube44\uc728', u'\uc555\ucd95 \ud6c4', u'\ucd5c\uc885 \uc808\uac10\ub960'],
    [
        [FORMAT_LABELS[f],
         '5,427 MB',
         f'{write_data[f]["ser_size"]/1024/1024:,.0f} MB',
         f'{write_data[f]["ser_ratio"]}%',
         f'{write_data[f]["comp_size"]/1024/1024:,.0f} MB',
         f'{write_data[f]["total_saving"]}%']
        for f in FORMATS
    ]
)
add_caption(doc, u'\ud45c 2. \ud3ec\ub9f7\ubcc4 \uc6a9\ub7c9 \ud6a8\uc728')

doc.add_paragraph()
doc.add_paragraph(
    u'Parquet\ub294 \ucee8\ub7fc\ud615 \uc778\ucf54\ub529\uc73c\ub85c \uc9c1\ub82c\ud654 \ub2e8\uacc4\uc5d0\uc11c \uc774\ubbf8 \uc6d0\ubcf8\uc758 30.4%\ub85c \uc904\uc5ec, '
    u'zstd \uc555\ucd95 \ud6c4 \ucd5c\uc885 96.2% \uc808\uac10\uc73c\ub85c \uac00\uc7a5 \ub192\uc740 \uc800\uc7a5 \ud6a8\uc728\uc744 \ub2ec\uc131\ud558\uc600\ub2e4. '
    u'Tar+JSON\uc740 JSON \ud14d\uc2a4\ud2b8\uc758 \uc7a5\ud669\uc131\uc73c\ub85c \uc9c1\ub82c\ud654 \uc2dc 125%\ub85c \ud655\ub300\ub418\uc9c0\ub9cc, '
    u'zstd\uac00 \ubc18\ubcf5\uc801\uc778 \ud14d\uc2a4\ud2b8 \ud328\ud134\uc744 \ud6a8\uacfc\uc801\uc73c\ub85c \uc555\ucd95\ud558\uc5ec(95.7%) \ucd5c\uc885 \uc6a9\ub7c9\uc740 \uacbd\uc7c1\uc801\uc774\ub2e4.'
)

doc.add_page_break()

# ============================================================
# 4. 읽기 벤치마크
# ============================================================
doc.add_heading(u'4. \uc77d\uae30 \ubca4\uce58\ub9c8\ud06c', level=1)
doc.add_paragraph(
    u'\uc77d\uae30 \ud30c\uc774\ud504\ub77c\uc778\uc740 .zst \ud30c\uc77c\uc744 \ud574\uc81c\ud55c \ud6c4 '
    u'\uc6d0\ubcf8 RegularLog \uac1d\uccb4\ub85c \uc5ed\uc9c1\ub82c\ud654\ud558\ub294 \uacfc\uc815\uc744 \uce21\uc815\ud55c\ub2e4. '
    u'\uae30\ubcf8 \uc131\ub2a5\uacfc \ud3ec\ub9f7\ubcc4 \ucd5c\uc801\ud654 \ud6a8\uacfc\ub97c \ud568\uaed8 \ud3c9\uac00\ud558\uae30 \uc704\ud574 '
    u'\ub2e8\uacc4\ubcc4\ub85c \ud14c\uc2a4\ud2b8\ub97c \uc218\ud589\ud558\uc600\ub2e4.'
)

doc.add_heading(u'\uc77d\uae30 \ub300\uc0c1 \ub370\uc774\ud130', level=3)
doc.add_paragraph(
    u'\ub85c\uadf8 \uc81c\ub108\ub808\uc774\ud130\ub85c \uc0dd\uc131\ud55c \uc6d0\ubcf8 .done \ud30c\uc77c\uc740 \ucd1d 5,000,000\uc904\uc774\ub2e4. '
    u'\uadf8\ub7ec\ub098 \uc77c\ubd80 \ub808\ucf54\ub4dc\uc758 original_log, attack_nm \ud544\ub4dc\uc5d0 \ud3ec\ud568\ub41c '
    u'\ud55c\uae00 \ud14d\uc2a4\ud2b8\uc5d0\uc11c \uac1c\ud589 \ubb38\uc790(\\n)\uac00 \ubc1c\uc0dd\ud558\uc5ec, '
    u'\ud558\ub098\uc758 JSON \ub808\ucf54\ub4dc\uac00 2~3\uc904\ub85c \ucaa4\uac1c\uc9c0\ub294 \ud604\uc0c1\uc774 \uc788\uc5c8\ub2e4. '
    u'LogReader\ub294 \uc904 \ub2e8\uc704\ub85c JSON\uc744 \ud30c\uc2f1\ud558\ubbc0\ub85c, '
    u'\ucaa4\uac1c\uc9c4 \uc904\uc740 \uc720\ud6a8\ud55c JSON\uc774 \uc544\ub2c8\uc5b4\uc11c \ud30c\uc2f1\uc5d0 \uc2e4\ud328\ud558\uace0 \uc6d0\ubcf8 \ub808\ucf54\ub4dc \uc790\uccb4\ub3c4 \ud568\uaed8 \uc81c\uc678\ub418\uc5c8\ub2e4.'
)
doc.add_paragraph(
    u'\uacb0\uacfc\uc801\uc73c\ub85c \uc4f0\uae30 \ubca4\uce58\ub9c8\ud06c\uc5d0\uc11c \uc2e4\uc81c\ub85c \uc9c1\ub82c\ud654\ub41c \ub808\ucf54\ub4dc\ub294 4,741,381\uac74\uc774\uba70, '
    u'\uc77d\uae30 \ubca4\uce58\ub9c8\ud06c\uc5d0\uc11c\ub3c4 \ub3d9\uc77c\ud558\uac8c 4,741,381\uac74\uc774 \ubcf5\uc6d0\ub418\uc5c8\ub2e4. '
    u'4\uac1c \ud3ec\ub9f7 \ubaa8\ub450 \ub3d9\uc77c\ud55c \uac74\uc218\ub97c \ucc98\ub9ac\ud558\uc600\uc73c\ubbc0\ub85c \ube44\uad50\uc758 \uacf5\uc815\uc131\uc5d0\ub294 \uc601\ud5a5\uc774 \uc5c6\ub2e4.'
)

add_table(doc,
    [u'\ud56d\ubaa9', u'\uac74\uc218'],
    [
        [u'\ub85c\uadf8 \uc81c\ub108\ub808\uc774\ud130 \uc0dd\uc131\ub7c9', '5,000,000'],
        [u'\uac1c\ud589 \ud3ec\ud568\uc73c\ub85c \uc778\ud55c \ud30c\uc2f1 \uc2e4\ud328', '258,619'],
        [u'\uc2e4\uc81c \ubca4\uce58\ub9c8\ud06c \ub300\uc0c1', '4,741,381'],
    ]
)
add_caption(doc, u'\ud45c 3-0. \uc77d\uae30 \ubca4\uce58\ub9c8\ud06c \ub300\uc0c1 \ub370\uc774\ud130 \uac74\uc218')

doc.add_paragraph()

# 4.1
doc.add_heading(u'4.1 \uc804\uccb4 \ucee8\ub7fc \uc77d\uae30 (Step 1)', level=2)
doc.add_picture(chart_paths['read_step1'], width=Inches(5.8))
add_caption(doc, u'\uadf8\ub9bc 3. \uc77d\uae30 \ubca4\uce58\ub9c8\ud06c Step 1 - \uc804\uccb4 \ucee8\ub7fc \uc5ed\uc9c1\ub82c\ud654 (5\ud68c \ud3c9\uade0)')

doc.add_paragraph()
add_table(doc,
    [u'\ud3ec\ub9f7', u'\ud574\uc81c(zstd)', u'\uc5ed\uc9c1\ub82c\ud654', u'\uc804\uccb4'],
    [
        [FORMAT_LABELS[f],
         f'{read_data[f]["step1"]["decomp"]:,} ms',
         f'{read_data[f]["step1"]["deser"]:,} ms',
         f'{read_data[f]["step1"]["total"]:,} ms']
        for f in FORMATS
    ]
)
add_caption(doc, u'\ud45c 3. \uc77d\uae30 \ubca4\uce58\ub9c8\ud06c Step 1 \ud3c9\uade0 \uc2dc\uac04 (ms)')

doc.add_paragraph()
doc.add_paragraph(
    u'Kryo\uac00 \ucd1d 45.7\ucd08\ub85c \uac00\uc7a5 \ube60\ub974\uba70, \ud2b9\ud788 \uc5ed\uc9c1\ub82c\ud654\uac00 6.9\ucd08\ub85c \uc555\ub3c4\uc801\uc774\ub2e4. '
    u'\uc2a4\ud0a4\ub9c8 \ud574\uc11d \uc5c6\uc774 POJO \ud544\ub4dc\uc5d0 \uc9c1\uc811 \ub9e4\ud551\ud558\ub294 \ubc14\uc774\ub108\ub9ac \uad6c\uc870 \ub355\ubd84\uc774\ub2e4. '
    u'Parquet\ub294 \ucd1d \uc2dc\uac04(50.1\ucd08)\uc5d0\uc11c \uacbd\uc7c1\ub825\uc774 \uc788\ub294\ub370, '
    u'\uc555\ucd95 \ud30c\uc77c\uc774 \uac00\uc7a5 \uc791\uc544(207 MB) \ud574\uc81c \uc2dc\uac04\uc774 \ube60\ub974\uae30 \ub54c\ubb38\uc774\ub2e4. '
    u'\ub2e4\ub9cc \ucee8\ub7fc-\ub85c\uc6b0 \ubcc0\ud658 \uacfc\uc815\uc73c\ub85c \uc5ed\uc9c1\ub82c\ud654\ub294 \ub290\ub9ac\ub2e4. '
    u'Tar+JSON\uc740 134.7\ucd08\ub85c \uac00\uc7a5 \ub290\ub9ac\ub2e4. JSON \ud14d\uc2a4\ud2b8 \ud30c\uc2f1\uc758 CPU \ubd80\ud558\uc640 '
    u'\ud070 \uc555\ucd95 \ud30c\uc77c(307 MB)\ub85c \uc778\ud55c \ud574\uc81c \uc2dc\uac04 \uc99d\uac00\uac00 \uc6d0\uc778\uc774\ub2e4.'
)

# 4.2 모듈별 최적화 읽기 비교
doc.add_heading(u'4.2 \ubaa8\ub4c8\ubcc4 \ucd5c\uc801\ud654 \uc77d\uae30 \ube44\uad50', level=2)
doc.add_paragraph(
    u'\ubaa8\ub4e0 \ud3ec\ub9f7\uc740 56\uac1c \ud544\ub4dc\uac00 \uc804\ubd80 \uc800\uc7a5\ub41c \ub3d9\uc77c\ud55c \ud30c\uc77c\uc744 \ub300\uc0c1\uc73c\ub85c, '
    u'\uc77d\ub294 \ucee8\ub7fc \uc218\ub97c 56(\uc804\uccb4) \u2192 40 \u2192 20 \u2192 10\uc73c\ub85c \uc904\uc5ec\uac00\uba70 \uc131\ub2a5\uc744 \uce21\uc815\ud558\uc600\ub2e4. '
    u'\uac01 \ud3ec\ub9f7\uc740 \uc790\uccb4 \ucd5c\uc801\ud654 \uae30\ubc95\uc744 \uc801\uc6a9\ud55c\ub2e4.'
)

doc.add_paragraph()
add_table(doc,
    [u'\ud3ec\ub9f7', u'\uc801\uc6a9 \ucd5c\uc801\ud654', u'\uc124\uba85'],
    [
        ['Tar+JSON', u'\uc5c6\uc74c', u'\ud589 \uc9c0\ud5a5 + JSON \ud14d\uc2a4\ud2b8 \uad6c\uc870\ub85c \ucd5c\uc801\ud654 \ubd88\uac00'],
        ['Kryo', u'POJO \uc7ac\uc0ac\uc6a9', u'\ub3d9\uc77c \uac1d\uccb4\uc5d0 \ud544\ub4dc \ub36e\uc5b4\uc4f0\uae30, GC \ubd80\ub2f4 \uac10\uc18c'],
        ['Parquet', u'\ucee8\ub7fc \ud504\ub8e8\ub2dd', u'\uc9c0\uc815\ub41c \ucee8\ub7fc\ub9cc \ub514\uc2a4\ud06c\uc5d0\uc11c \uc77d\uae30 (I/O \uc808\uac10)'],
        ['Avro', 'SpecificDatumReader', u'\ucf54\ub4dc\uc83c \ud074\ub798\uc2a4\ub85c \uc9c1\uc811 \uc811\uadfc (\ub9ac\ud50c\ub809\uc158 \uc81c\uac70)'],
    ]
)
add_caption(doc, u'\ud45c 3-1. \ud3ec\ub9f7\ubcc4 \uc801\uc6a9 \ucd5c\uc801\ud654 \uae30\ubc95')

doc.add_paragraph()
doc.add_picture(chart_paths['optimized_read'], width=Inches(5.8))
add_caption(doc, u'\uadf8\ub9bc 4. \ubaa8\ub4c8\ubcc4 \ucd5c\uc801\ud654 \uc804\ud6c4 \uc5ed\uc9c1\ub82c\ud654 \uc2dc\uac04 \ube44\uad50')

doc.add_paragraph()

# 표 3-2: 세로축=포맷, 가로축=읽는 컬럼 수 + 개선율
def _imp(base, opt):
    if base == 0:
        return '-'
    pct = (1 - opt / base) * 100
    if abs(pct) < 1:
        return '-'
    return f'-{pct:.1f}%'

add_table(doc,
    [u'\ud3ec\ub9f7',
     u'56col\n(\uc804\uccb4)',
     u'40col', u'\uac1c\uc120\uc728',
     u'20col', u'\uac1c\uc120\uc728',
     u'10col', u'\uac1c\uc120\uc728'],
    [
        ['Tar+JSON',
         f'{read_data["tar"]["step1"]["deser"]:,}',
         f'{read_data["tar"]["step2"]["deser"]:,}', _imp(read_data["tar"]["step1"]["deser"], read_data["tar"]["step2"]["deser"]),
         f'{read_data["tar"]["step2"]["deser"]:,}', '-',
         f'{read_data["tar"]["step2"]["deser"]:,}', '-'],
        ['Kryo',
         f'{read_data["kryo"]["step1"]["deser"]:,}',
         f'{read_data["kryo"]["step2"]["deser"]:,}', _imp(read_data["kryo"]["step1"]["deser"], read_data["kryo"]["step2"]["deser"]),
         f'{read_data["kryo"]["step2"]["deser"]:,}', '-',
         f'{read_data["kryo"]["step2"]["deser"]:,}', '-'],
        ['Parquet',
         f'{read_data["parquet"]["step1"]["deser"]:,}',
         f'{read_data["parquet"]["step2"]["deser"]:,}', _imp(read_data["parquet"]["step1"]["deser"], read_data["parquet"]["step2"]["deser"]),
         f'{read_data["parquet"]["step3"]["deser"]:,}', _imp(read_data["parquet"]["step1"]["deser"], read_data["parquet"]["step3"]["deser"]),
         f'{read_data["parquet"]["step4"]["deser"]:,}', _imp(read_data["parquet"]["step1"]["deser"], read_data["parquet"]["step4"]["deser"])],
        ['Avro',
         f'{read_data["avro"]["step1"]["deser"]:,}',
         f'{read_data["avro"]["step2"]["deser"]:,}', _imp(read_data["avro"]["step1"]["deser"], read_data["avro"]["step2"]["deser"]),
         f'{read_data["avro"]["step2"]["deser"]:,}', '-',
         f'{read_data["avro"]["step2"]["deser"]:,}', '-'],
    ]
)
add_caption(doc, u'\ud45c 3-2. \uc5ed\uc9c1\ub82c\ud654 \uc2dc\uac04(ms) - \uc804\uccb4 56\uac1c \ud544\ub4dc \uc800\uc7a5 \ub300\uc0c1, \uc77d\ub294 \ucee8\ub7fc \uc218\ubcc4 \ube44\uad50')

doc.add_paragraph()
doc.add_paragraph(
    u'Parquet\ub9cc\uc774 \ucee8\ub7fc \uc218\uac00 \uc904\uc5b4\ub4e4\uc218\ub85d \uc5ed\uc9c1\ub82c\ud654 \uc2dc\uac04\uc774 \uc120\ud615\uc801\uc73c\ub85c \uac10\uc18c\ud55c\ub2e4. '
    u'\uc774\ub294 \ucee8\ub7fc\ud615 \uc800\uc7a5\uc758 \uadfc\ubcf8\uc801\uc778 \uac15\uc810\uc73c\ub85c, \ud544\uc694\ud55c \ucee8\ub7fc\ub9cc \ub514\uc2a4\ud06c\uc5d0\uc11c \uc77d\uae30 \ub54c\ubb38\uc774\ub2e4. '
    u'\ubc18\uba74 Tar+JSON, Kryo, Avro\ub294 \ud589 \uc9c0\ud5a5 \ud3ec\ub9f7\uc774\ubbc0\ub85c \ucee8\ub7fc \uc218\ub97c \uc904\uc5ec\ub3c4 '
    u'\uc804\uccb4 \ub370\uc774\ud130\ub97c \uc77d\uc5b4\uc57c \ud558\uba70, \uac01\uc790\uc758 \ucd5c\uc801\ud654(POJO \uc7ac\uc0ac\uc6a9, SpecificDatumReader)\ub97c '
    u'\uc801\uc6a9\ud574\ub3c4 \ucee8\ub7fc \uc218 \ubcc0\ud654\uc640\ub294 \ubb34\uad00\ud558\uac8c \ub3d9\uc77c\ud55c \uc131\ub2a5\uc744 \ubcf4\uc778\ub2e4.'
)

doc.add_paragraph()
doc.add_paragraph(
    u'\uc5ed\uc9c1\ub82c\ud654 \uc2dc\uac04\uc740 \ucee8\ub7fc \uc218\uc5d0 \uac70\uc758 \uc120\ud615\uc801\uc73c\ub85c \ube44\ub840\ud55c\ub2e4. '
    u'10\uac1c \ucee8\ub7fc(\uc804\uccb4\uc758 18%)\ub9cc \uc77d\uc744 \uacbd\uc6b0 \uc5ed\uc9c1\ub82c\ud654\uac00 4.6\ucd08\ub85c, '
    u'\uc804\uccb4 \ucee8\ub7fc \ub300\ube44 83% \uac1c\uc120\ub41c\ub2e4. '
    u'\uc774\ub294 \uc77c\ubd80 \ud544\ub4dc\ub9cc \uc870\ud68c\ud558\ub294 \ubd84\uc11d/\ubaa8\ub2c8\ud130\ub9c1 \ucffc\ub9ac\uc5d0 \uc9c1\uc811\uc801\uc778 \uc774\uc810\uc744 \uc81c\uacf5\ud55c\ub2e4. '
    u'\ucc38\uace0\ub85c \ud574\uc81c \uc2dc\uac04\uc740 .zst \ud30c\uc77c \uc804\uccb4\ub97c \ud480\uc5b4\uc57c \ud558\ubbc0\ub85c \ucee8\ub7fc \uc218\uc640 \ubb34\uad00\ud558\uac8c \uc77c\uc815\ud558\ub2e4.'
)

doc.add_page_break()

# ============================================================
# 5. 쓰기/읽기 종합 비교
# ============================================================
doc.add_heading(u'5. \uc4f0\uae30/\uc77d\uae30 \uc885\ud569 \ube44\uad50', level=1)
doc.add_picture(chart_paths['total_rw'], width=Inches(5.8))
add_caption(doc, u'\uadf8\ub9bc 6. \uc4f0\uae30 vs \uc77d\uae30 \uc804\uccb4 \uc2dc\uac04 \ube44\uad50')

doc.add_paragraph()
add_table(doc,
    [u'\ud3ec\ub9f7', u'\uc4f0\uae30 \uc804\uccb4', u'\uc77d\uae30 \uc804\uccb4', u'\uc4f0\uae30/\uc77d\uae30 \ube44\uc728'],
    [
        [FORMAT_LABELS[f],
         f'{write_data[f]["total"]:,} ms',
         f'{read_data[f]["step1"]["total"]:,} ms',
         f'{write_data[f]["total"] / read_data[f]["step1"]["total"]:.1f}x']
        for f in FORMATS
    ]
)
add_caption(doc, u'\ud45c 5. \ud3ec\ub9f7\ubcc4 \uc4f0\uae30/\uc77d\uae30 \uc2dc\uac04 \ube44\uc728')

doc.add_paragraph()
doc.add_paragraph(
    u'\ubaa8\ub4e0 \ud3ec\ub9f7\uc774 \uc4f0\uae30 > \uc77d\uae30 \ud2b9\uc131\uc744 \ubcf4\uc778\ub2e4. '
    u'\uc4f0\uae30\uc5d0\ub294 \ud30c\uc2f1\uacfc \uc9c1\ub82c\ud654\uac00 \ubaa8\ub450 \ud3ec\ud568\ub418\ub294 \ubc18\uba74 '
    u'\uc77d\uae30\ub294 \uc5ed\uc9c1\ub82c\ud654\ub9cc \ud544\uc694\ud558\uae30 \ub54c\ubb38\uc774\ub2e4. '
    u'Kryo\uac00 \uac00\uc7a5 \uade0\ud615 \uc7a1\ud78c \ud504\ub85c\ud30c\uc77c(1.7x)\uc744 \ubcf4\uc774\uba70, '
    u'\uc77d\uae30/\uc4f0\uae30\uac00 \ubaa8\ub450 \ube48\ubc88\ud55c \ud30c\uc774\ud504\ub77c\uc778\uc5d0 \uc801\ud569\ud558\ub2e4.'
)

doc.add_page_break()

# ============================================================
# 6. 발생 이슈 및 조치 사항
# ============================================================
doc.add_heading(u'6. \ubc1c\uc0dd \uc774\uc288 \ubc0f \uc870\uce58 \uc0ac\ud56d', level=1)
doc.add_paragraph(
    u'\ubca4\uce58\ub9c8\ud06c \uac1c\ubc1c \uacfc\uc815\uc5d0\uc11c \ubc1c\uc0dd\ud55c \uae30\uc220\uc801 \uc774\uc288\uc640 \ud574\uacb0 \uacfc\uc815\uc744 \uae30\ub85d\ud55c\ub2e4. '
    u'\ubaa8\ub4e0 \uc870\uce58\ub294 4\uac1c \ud3ec\ub9f7 \uac04 \uce21\uc815 \uacf5\uc815\uc131\uc744 \uc720\uc9c0\ud558\ub294 \ubc29\ud5a5\uc73c\ub85c \uc124\uacc4\ub418\uc5c8\ub2e4.'
)

issues = [
    {
        'title': u'6.1 OutOfMemoryError - ByteArrayOutputStream (2GB \ubc30\uc5f4 \ud55c\uacc4)',
        'problem': u'\ucd08\uae30 \uad6c\ud604\uc5d0\uc11c ByteArrayOutputStream\uc73c\ub85c \uc9c1\ub82c\ud654 \ub370\uc774\ud130\ub97c \uba54\ubaa8\ub9ac\uc5d0 \ucd95\uc801\ud558\uc600\ub2e4. '
                   u'1,086\uac1c \ud30c\uc77c(5.3 GB) \ucc98\ub9ac \uc2dc Java\uc758 ~2 GB \ubc30\uc5f4 \ud06c\uae30 \ud55c\uacc4\ub97c \ucd08\uacfc\ud558\uc600\uace0, '
                   u'toByteArray() \ud638\ucd9c \uc2dc \uba54\ubaa8\ub9ac\uac00 2\ubc30\ub85c \uc99d\uac00\ud558\uc5ec(\ub370\uc774\ud130 + \ubcf5\uc0ac\ubcf8 = ~6.6 GB vs 4 GB \ud799) OOM\uc774 \ubc1c\uc0dd\ud558\uc600\ub2e4.',
        'fix': u'\uba54\ubaa8\ub9ac \ubc84\ud37c \ub300\uc2e0 \uc784\uc2dc \ud30c\uc77c \uc4f0\uae30(FileOutputStream + BufferedOutputStream)\ub85c \uc804\ud658\ud558\uc600\ub2e4. '
               u'4\uac1c \uc778\ucf54\ub354 \ubaa8\ub450\uc5d0 \ub3d9\uc77c\ud558\uac8c \uc801\uc6a9\ud558\uc600\uc73c\uba70, \uc784\uc2dc \ud30c\uc77c\uc740 \uc2a4\ud2b8\ub9bc close \uc2dc \uc790\ub3d9 \uc0ad\uc81c\ub41c\ub2e4. '
               u'OS \ud398\uc774\uc9c0 \uce90\uc2dc\uac00 \ub514\uc2a4\ud06c I/O \ubd80\ub2f4\uc744 \uc644\ud654\ud55c\ub2e4.',
    },
    {
        'title': u'6.2 OutOfMemoryError - \uc804\uccb4 \ub85c\ub529 \uba54\ubaa8\ub9ac \ubd80\uc871 (~26 GB)',
        'problem': u'\uc57d 500\ub9cc \uac74\uc758 RegularLog \uac1d\uccb4\ub97c List\uc5d0 \uc804\uccb4 \ub85c\ub529\ud55c \ud6c4 \uc9c1\ub82c\ud654\ud558\ub294 \ubc29\uc2dd\uc73c\ub85c, '
                   u'\uc57d 26 GB \uba54\ubaa8\ub9ac\uac00 \ud544\uc694\ud558\uc5ec 4 GB \ud799\uc744 \ucd08\uacfc\ud558\uc600\ub2e4.',
        'fix': u'\ud30c\uc77c \ub2e8\uc704 \uc2a4\ud2b8\ub9ac\ubc0d \ucc98\ub9ac\ub85c \uc804\ud658\ud558\uc600\ub2e4. 1,086\uac1c \ud30c\uc77c(\ud30c\uc77c\ub2f9 \uc57d 4,600\uac74)\uc744 \ud558\ub098\uc529 \uc77d\uace0, '
               u'\ud30c\uc2f1\ud558\uace0, \uc9c1\ub82c\ud654\ud558\uba70, nanoTime()\uc73c\ub85c \uac01 \ub2e8\uacc4 \uc2dc\uac04\uc744 \ub204\uc801 \uce21\uc815\ud55c\ub2e4. '
               u'\ud53c\ud06c \uba54\ubaa8\ub9ac\uac00 \uc57d 25 MB(\ud799\uc758 0.6%)\ub85c \uac10\uc18c\ud558\uc600\uc73c\uba70, 4\uac1c \ud3ec\ub9f7 \ubaa8\ub450 \ub3d9\uc77c\ud558\uac8c \uc801\uc6a9\ud558\uc600\ub2e4.',
    },
    {
        'title': u'6.3 Tar+JSON \ubca4\uce58\ub9c8\ud06c \uacf5\uc815\uc131 - StringBuilder OOM',
        'problem': u'TarMain\uc5d0\uc11c StringBuilder\ub85c \ubaa8\ub4e0 JSON \ubb38\uc790\uc5f4\uc744 \uba54\ubaa8\ub9ac\uc5d0 \uc5f0\uacb0\ud558\uc5ec '
                   u'500\ub9cc \uac74 \uae30\uc900 \uc57d 7 GB\uac00 \ud544\uc694\ud558\uc600\ub2e4. OOM \ubc1c\uc0dd \ubc0f \uacfc\ub3c4\ud55c GC\ub85c \uc778\ud574 '
                   u'tar\uc758 \uc9c1\ub82c\ud654 \uc2dc\uac04\uc774 \ubd88\uacf5\uc815\ud558\uac8c \ubd80\ud480\ub824\uc84c\ub2e4.',
        'fix': u'\uc21c\ucc28\uc801 \ub514\uc2a4\ud06c \uc4f0\uae30(POJO \u2192 JSON bytes \u2192 BufferedOutputStream \u2192 \uc784\uc2dc \ud30c\uc77c)\ub85c \uc804\ud658\ud558\uc5ec '
               u'Kryo, Parquet, Avro\uc640 \ub3d9\uc77c\ud55c I/O \ud328\ud134\uc744 \uc801\uc6a9\ud558\uc600\ub2e4. OOM \uc704\ud5d8\uacfc GC \ud3b8\ud5a5\uc744 \uc81c\uac70\ud558\uc600\ub2e4.',
    },
    {
        'title': u'6.4 Avro NullPointerException - Null \ud544\ub4dc \ucc98\ub9ac',
        'problem': u'Avro\uc758 \uae30\ubcf8 \uc2a4\ud0a4\ub9c8\ub294 non-null String\uc744 \uac15\uc81c\ud55c\ub2e4. '
                   u'RegularLog \ud544\ub4dc\uc5d0 null \uac12\uc774 \ud3ec\ud568\ub418\uba74 NullPointerException\uc774 \ubc1c\uc0dd\ud558\uc600\ub2e4. '
                   u'\ub2e4\ub978 \ud3ec\ub9f7(Kryo, Parquet, Tar)\uc740 null\uc744 \uc790\uccb4\uc801\uc73c\ub85c \ucc98\ub9ac\ud55c\ub2e4.',
        'fix': u'LogReader.fillNulls() \uc804\ucc98\ub9ac\ub97c \ucd94\uac00\ud558\uc5ec \ubaa8\ub4e0 null String \ud544\ub4dc\ub97c \ube48 \ubb38\uc790\uc5f4("")\ub85c \ub300\uccb4\ud558\uc600\ub2e4. '
               u'\uc778\ucf54\ub354 \uc804\ub2e8\uacc4\uc5d0 \uc804\uc5ed\uc801\uc73c\ub85c \uc801\uc6a9\ud558\uc5ec 4\uac1c \ud3ec\ub9f7\uc774 \ub3d9\uc77c\ud55c \ub370\uc774\ud130\ub97c \ucc98\ub9ac\ud558\ub3c4\ub85d \ud558\uc600\ub2e4.',
    },
    {
        'title': u'6.5 Parquet \uc77d\uae30 ClassCastException - \uc2a4\ud0a4\ub9c8 \ub124\uc784\uc2a4\ud398\uc774\uc2a4 \ubd88\uc77c\uce58',
        'problem': u'Parquet \uc77d\uae30 \ubca4\uce58\ub9c8\ud06c\uc5d0\uc11c ClassCastException\uc774 \ubc1c\uc0dd\ud558\uc600\ub2e4. '
                   u'GenericData$Record\ub97c RegularLog\ub85c \uce90\uc2a4\ud305\ud560 \uc218 \uc5c6\ub294 \uc624\ub958\ub85c, '
                   u'\uadfc\ubcf8 \uc6d0\uc778\uc740 Parquet \ud30c\uc77c\uc5d0 \ub0b4\uc7a5\ub41c \uc4f0\uae30 \uc2a4\ud0a4\ub9c8\uc758 \ub124\uc784\uc2a4\ud398\uc774\uc2a4(com.vortexlake.bench.model)\uc640 '
                   u'\uc77d\uae30 POJO\uc758 \ub124\uc784\uc2a4\ud398\uc774\uc2a4(com.vortexlake.readbench.model)\uac00 \ub2ec\ub790\uae30 \ub54c\ubb38\uc774\ub2e4. '
                   u'ReflectData\uac00 \uc4f0\uae30 \uc2a4\ud0a4\ub9c8\uc758 \ud074\ub798\uc2a4\ub97c \ucc3e\uc9c0 \ubabb\ud574 Generic \ubaa8\ub4dc\ub85c \uc790\ub3d9 \uc804\ud658\ub418\uc5c8\ub2e4.',
        'fix': u'AvroReadSupport.setAvroReadSchema()\ub85c \uc77d\uae30 \uc2a4\ud0a4\ub9c8\ub97c \uba85\uc2dc\uc801\uc73c\ub85c \uc9c0\uc815\ud558\uace0, '
               u'conf.setClass()\ub85c ReflectDataSupplier\ub97c \uac15\uc81c \uc124\uc815\ud558\uc5ec Reflect \ubaa8\ub4dc\ub97c \uc720\uc9c0\ud558\uc600\ub2e4. '
               u'\ucee8\ub7fc \ud504\ub8e8\ub2dd(Step 2~4)\uc5d0\ub294 AvroReadSupport.setRequestedProjection()\uc744 \uc0ac\uc6a9\ud558\uc600\ub2e4.',
    },
]

for issue in issues:
    doc.add_heading(issue['title'], level=2)
    p = doc.add_paragraph()
    run = p.add_run(u'\ud604\uc0c1: ')
    run.bold = True
    p.add_run(issue['problem'])
    p = doc.add_paragraph()
    run = p.add_run(u'\uc870\uce58: ')
    run.bold = True
    p.add_run(issue['fix'])
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 7. 결론
# ============================================================
doc.add_heading(u'7. \uacb0\ub860', level=1)

doc.add_paragraph(
    u'\uc4f0\uae30/\uc77d\uae30 \ubca4\uce58\ub9c8\ud06c \uacb0\uacfc\ub97c \uc885\ud569\ud558\uba74, '
    u'\uac01 \ud3ec\ub9f7\uc740 \uc6b4\uc601 \uc6b0\uc120\uc21c\uc704\uc5d0 \ub530\ub77c \ub2e4\ub978 \uac15\uc810\uc744 \ubcf4\uc778\ub2e4.'
)

conclusions = [
    ['Kryo',
     u'\uac00\uc7a5 \ube60\ub978 \uc4f0\uae30(77.6\ucd08)/\uc77d\uae30(45.7\ucd08) \uc131\ub2a5. '
     u'\ucf64\ud329\ud2b8 \ubc14\uc774\ub108\ub9ac \ud3ec\ub9f7\uc73c\ub85c \uc9c1\ub82c\ud654 \ud6c4 \uc6d0\ubcf8\uc758 49.9%. '
     u'\uace0\uc18d \ucc98\ub9ac\uac00 \ud544\uc694\ud55c \ud30c\uc774\ud504\ub77c\uc778\uc5d0 \uc801\ud569. '
     u'\uc81c\ud55c: \ucee8\ub7fc \ud504\ub8e8\ub2dd \ubd88\uac00, \uc2a4\ud0a4\ub9c8 \uc9c4\ud654 \uc2dc \uc8fc\uc758 \ud544\uc694.'],
    ['Parquet',
     u'\ucd5c\uace0 \uc800\uc7a5 \ud6a8\uc728(96.2% \uc808\uac10) \ubc0f \ucee8\ub7fc \ud504\ub8e8\ub2dd \uc9c0\uc6d0. '
     u'10\uac1c \ucee8\ub7fc \uc77d\uae30 \uc2dc \uc5ed\uc9c1\ub82c\ud654 83% \uac1c\uc120. '
     u'\uc4f0\uae30 \uc18d\ub3c4\ub294 \ubcf4\ud1b5(124.5\ucd08). '
     u'\ubd84\uc11d \uc911\uc2ec \uc6cc\ud06c\ub85c\ub4dc\uc5d0 \ucd5c\uc801.'],
    ['Avro',
     u'\uade0\ud615 \uc7a1\ud78c \uc131\ub2a5. \uc4f0\uae30 85.9\ucd08, \uc77d\uae30 63.6\ucd08. '
     u'SpecificDatumReader \ucd5c\uc801\ud654 \uac00\ub2a5. '
     u'\uc2a4\ud0a4\ub9c8 \uc9c4\ud654 \ubc0f \ud06c\ub85c\uc2a4 \uc5b8\uc5b4 \ud638\ud658\uc131\uc774 \uc911\uc694\ud55c \uacbd\uc6b0 \uc801\ud569.'],
    ['Tar+JSON',
     u'\uc4f0\uae30(181.0\ucd08)/\uc77d\uae30(134.7\ucd08) \ubaa8\ub450 \uac00\uc7a5 \ub290\ub9bc. '
     u'JSON \ud14d\uc2a4\ud2b8\uc758 \ubc18\ubcf5 \ud328\ud134\uc73c\ub85c zstd \uc555\ucd95\ub960\uc740 \uc6b0\uc218(95.7%). '
     u'\uc0ac\ub78c\uc774 \uc77d\uc744 \uc218 \uc788\ub294 \uc911\uac04 \ucd9c\ub825\ubb3c \uc0dd\uc131. '
     u'\uc544\uce74\uc774\ube0c/\ub514\ubc84\uae45 \uc6a9\ub3c4\uc5d0 \uc801\ud569.'],
]

add_table(doc,
    [u'\ud3ec\ub9f7', u'\ud3c9\uac00'],
    conclusions
)
add_caption(doc, u'\ud45c 6. \ud3ec\ub9f7\ubcc4 \uc801\ud569\uc131 \uc694\uc57d')

doc.add_paragraph()
doc.add_heading(u'\uad8c\uace0 \uc804\ub7b5', level=2)
doc.add_paragraph(
    u'VortexLake \ub85c\uadf8 \uc800\uc7a5 \ud30c\uc774\ud504\ub77c\uc778\uc5d0\ub294 \uc774\uc911 \ud3ec\ub9f7 \uc804\ub7b5\uc744 \uad8c\uace0\ud55c\ub2e4. '
    u'\uc2e4\uc2dc\uac04 \uc218\uc9d1 \uacbd\ub85c(Hot Path)\uc5d0\ub294 Kryo\ub97c \uc801\uc6a9\ud558\uc5ec \ucd5c\ub300 \ucc98\ub9ac\ub7c9\uc744 \ud655\ubcf4\ud558\uace0, '
    u'\uc7a5\uae30 \ubd84\uc11d \uc800\uc7a5\uc18c(Cold/Analytics Path)\uc5d0\ub294 Parquet\ub97c \uc801\uc6a9\ud558\uc5ec '
    u'\ucee8\ub7fc \ud504\ub8e8\ub2dd\uc73c\ub85c \ud558\uc704 \ucffc\ub9ac \uc131\ub2a5\uc744 \uadf9\ub300\ud654\ud558\ub294 \ubc29\uc548\uc774\ub2e4.'
)

doc.add_page_break()

# ============================================================
# 부록 A. RegularLog 필드 명세
# ============================================================
doc.add_heading(u'\ubd80\ub85d A. RegularLog \ud544\ub4dc \uba85\uc138', level=1)
doc.add_paragraph(
    u'RegularLog\ub294 \ubcf4\uc548 \ub85c\uadf8\ub97c \uc815\uaddc\ud654\ud55c POJO \ud074\ub798\uc2a4\ub85c, '
    u'\ucd1d 56\uac1c \ud544\ub4dc\ub85c \uad6c\uc131\ub41c\ub2e4. \ubaa8\ub4e0 \ud544\ub4dc\ub294 String \ud0c0\uc785\uc774\uba70, '
    u'\uae30\ub2a5\ubcc4\ub85c 8\uac1c \uadf8\ub8f9\uc73c\ub85c \ubd84\ub958\ub41c\ub2e4.'
)

# 전체 56개 필드 (4열 테이블)
all_fields = [
    (1,  'eqpIp',             u'\uc7a5\ube44 IP'),
    (2,  'eqpDt',             u'\uc7a5\ube44 \uc77c\uc2dc'),
    (3,  'recvTime',          u'\uc218\uc2e0 \uc2dc\uac04'),
    (4,  'agentIp',           u'\uc5d0\uc774\uc804\ud2b8 IP'),
    (5,  'connectIp',         u'\uc811\uc18d IP'),
    (6,  'macAddr',           u'MAC \uc8fc\uc18c'),
    (7,  'logType',           u'\ub85c\uadf8 \uc720\ud615'),
    (8,  'dataType',          u'\ub370\uc774\ud130 \uc720\ud615'),
    (9,  'normalizeResult',   u'\uc815\uaddc\ud654 \uacb0\uacfc'),
    (10, 'key',               u'\ud0a4'),
    (11, 'parserKey',         u'\ud30c\uc11c \ud0a4'),
    (12, 'eqpDtYear',         u'\uc5f0\ub3c4'),
    (13, 'eqpDtMonth',        u'\uc6d4'),
    (14, 'eqpDtDate',         u'\ub0a0\uc9dc'),
    (15, 'eqpDtDayOfMonth',   u'\uc77c'),
    (16, 'eqpDtDayOfWeek',    u'\uc694\uc77c'),
    (17, 'eqpDtHour',         u'\uc2dc\uac04'),
    (18, 'eqpDtTime',         u'\uc2dc\uac01'),
    (19, 'eqpDtHolidayYn',    u'\uacf5\ud734\uc77c \uc5ec\ubd80'),
    (20, 'startTime',         u'\uc2dc\uc791 \uc2dc\uac04'),
    (21, 'endTime',           u'\uc885\ub8cc \uc2dc\uac04'),
    (22, 'eventTime',         u'\uc774\ubca4\ud2b8 \uc2dc\uac04'),
    (23, 'gatherEtime',       u'\uc218\uc9d1 \uc885\ub8cc \uc2dc\uac04'),
    (24, 'srcIp',             u'\ucd9c\ubc1c\uc9c0 IP'),
    (25, 'srcPort',           u'\ucd9c\ubc1c\uc9c0 \ud3ec\ud2b8'),
    (26, 'srcCountryCode',    u'\ucd9c\ubc1c\uc9c0 \uad6d\uac00\ucf54\ub4dc'),
    (27, 'srcCountryName',    u'\ucd9c\ubc1c\uc9c0 \uad6d\uac00\uba85'),
    (28, 'srcBlackip',        u'\ucd9c\ubc1c\uc9c0 \ube14\ub799IP'),
    (29, 'srcBlackipSeverity',u'\ucd9c\ubc1c\uc9c0 \ube14\ub799IP \uc2ec\uac01\ub3c4'),
    (30, 'srcAssetGroupCd',   u'\ucd9c\ubc1c\uc9c0 \uc790\uc0b0\uadf8\ub8f9'),
    (31, 'srcAssetId',        u'\ucd9c\ubc1c\uc9c0 \uc790\uc0b0ID'),
    (32, 'srcAssetNm',        u'\ucd9c\ubc1c\uc9c0 \uc790\uc0b0\uba85'),
    (33, 'dstnIp',            u'\ubaa9\uc801\uc9c0 IP'),
    (34, 'dstnPort',          u'\ubaa9\uc801\uc9c0 \ud3ec\ud2b8'),
    (35, 'dstnCountryCode',   u'\ubaa9\uc801\uc9c0 \uad6d\uac00\ucf54\ub4dc'),
    (36, 'dstnCountryName',   u'\ubaa9\uc801\uc9c0 \uad6d\uac00\uba85'),
    (37, 'dstnBlackip',       u'\ubaa9\uc801\uc9c0 \ube14\ub799IP'),
    (38, 'dstnBlackipSeverity',u'\ubaa9\uc801\uc9c0 \ube14\ub799IP \uc2ec\uac01\ub3c4'),
    (39, 'dstnAssetGroupCd',  u'\ubaa9\uc801\uc9c0 \uc790\uc0b0\uadf8\ub8f9'),
    (40, 'dstnAssetId',       u'\ubaa9\uc801\uc9c0 \uc790\uc0b0ID'),
    (41, 'dstnAssetNm',       u'\ubaa9\uc801\uc9c0 \uc790\uc0b0\uba85'),
    (42, 'prtc',              u'\ud504\ub85c\ud1a0\ucf5c'),
    (43, 'msg',               u'\uba54\uc2dc\uc9c0'),
    (44, 'attackNm',          u'\uacf5\uaca9\uba85'),
    (45, 'action',            u'\uc561\uc158'),
    (46, 'faultMsg',          u'\uc7a5\uc560 \uba54\uc2dc\uc9c0'),
    (47, 'inOut',             u'\ub0b4/\uc678\ubd80'),
    (48, 'userId',            u'\uc0ac\uc6a9\uc790 ID'),
    (49, 'instCd1',           u'\uae30\uad00\ucf54\ub4dc 1'),
    (50, 'instCd2',           u'\uae30\uad00\ucf54\ub4dc 2'),
    (51, 'cnt',               u'\uce74\uc6b4\ud2b8'),
    (52, 'count',             u'\uac74\uc218'),
    (53, 'originalLog',       u'\uc6d0\ubcf8 \ub85c\uadf8'),
    (54, 'parseFailMsg',      u'\ud30c\uc2f1 \uc2e4\ud328 \uba54\uc2dc\uc9c0'),
    (55, 'test',              u'\ud14c\uc2a4\ud2b8'),
    (56, 'test1',             u'\ud14c\uc2a4\ud2b8 1'),
]

# 그룹 정보
groups = [
    (u'\uc7a5\ube44/\uc218\uc9d1 \uae30\ubcf8 \uc815\ubcf4', 1, 11),
    (u'\ub0a0\uc9dc/\uc2dc\uac04 \ubd84\ud574', 12, 19),
    (u'\uc774\ubca4\ud2b8 \uc2dc\uac04', 20, 23),
    (u'\ucd9c\ubc1c\uc9c0 \uc815\ubcf4', 24, 32),
    (u'\ubaa9\uc801\uc9c0 \uc815\ubcf4', 33, 41),
    (u'\ud504\ub85c\ud1a0\ucf5c/\uc774\ubca4\ud2b8', 42, 48),
    (u'\uae30\uad00/\uce74\uc6b4\ud2b8', 49, 52),
    (u'\uc6d0\ubcf8/\ud14c\uc2a4\ud2b8', 53, 56),
]

# 전체 필드 테이블 (4열: 번호, 필드명, 설명, 그룹)
doc.add_heading(u'A.1 \uc804\uccb4 \ud544\ub4dc \ubaa9\ub85d (56\uac1c)', level=2)

# 3열 x 2세트 = 6열 테이블로 한 페이지에 맞추기
half = 28
table = doc.add_table(rows=1 + half, cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 헤더
for col_offset in [0, 3]:
    for i, h in enumerate([u'#', u'\ud544\ub4dc\uba85', u'\uc124\uba85']):
        cell = table.rows[0].cells[col_offset + i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

# 데이터
for row_idx in range(half):
    # 왼쪽 열 (1~28)
    idx_l = row_idx
    cell_num = table.rows[row_idx + 1].cells[0]
    cell_name = table.rows[row_idx + 1].cells[1]
    cell_desc = table.rows[row_idx + 1].cells[2]
    cell_num.text = str(all_fields[idx_l][0])
    cell_name.text = all_fields[idx_l][1]
    cell_desc.text = all_fields[idx_l][2]
    for c in [cell_num, cell_name, cell_desc]:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

    # 오른쪽 열 (29~56)
    idx_r = row_idx + half
    cell_num2 = table.rows[row_idx + 1].cells[3]
    cell_name2 = table.rows[row_idx + 1].cells[4]
    cell_desc2 = table.rows[row_idx + 1].cells[5]
    cell_num2.text = str(all_fields[idx_r][0])
    cell_name2.text = all_fields[idx_r][1]
    cell_desc2.text = all_fields[idx_r][2]
    for c in [cell_num2, cell_name2, cell_desc2]:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

add_caption(doc, u'\ud45c A-1. RegularLog \uc804\uccb4 56\uac1c \ud544\ub4dc \ubaa9\ub85d')

doc.add_paragraph()

# 그룹 요약 테이블
grp_table = doc.add_table(rows=1 + len(groups), cols=3)
grp_table.style = 'Light Grid Accent 1'
grp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate([u'\uadf8\ub8f9', u'\ud544\ub4dc \ubc94\uc704', u'\ud544\ub4dc \uc218']):
    cell = grp_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
for r_idx, (name, start, end) in enumerate(groups):
    grp_table.rows[r_idx + 1].cells[0].text = name
    grp_table.rows[r_idx + 1].cells[1].text = f'#{start} ~ #{end}'
    grp_table.rows[r_idx + 1].cells[2].text = str(end - start + 1)
    for c_idx in range(3):
        for p in grp_table.rows[r_idx + 1].cells[c_idx].paragraphs:
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)

add_caption(doc, u'\ud45c A-2. \ud544\ub4dc \uadf8\ub8f9 \ubd84\ub958')

doc.add_page_break()

# ============================================================
# 부록 B. 컬럼 프루닝 필드 명세
# ============================================================
doc.add_heading(u'\ubd80\ub85d B. \ucee8\ub7fc \ud504\ub8e8\ub2dd \ud544\ub4dc \uba85\uc138', level=1)
doc.add_paragraph(
    u'Parquet \uc77d\uae30 \ubca4\uce58\ub9c8\ud06c\uc758 Step 2~4\uc5d0\uc11c\ub294 ALL_FIELDS \ubc30\uc5f4 \uc21c\uc11c\ub300\ub85c '
    u'\uc55e\uc5d0\uc11c\ubd80\ud130 40/20/10\uac1c\ub97c \uc120\ud0dd\ud558\uc5ec \ucee8\ub7fc \ud504\ub8e8\ub2dd\uc744 \uc218\ud589\ud55c\ub2e4. '
    u'\uc774 \ubc29\uc2dd\uc740 \uc2e4\uc81c \uc6b4\uc601\uc5d0\uc11c \uae30\ubcf8 \uc815\ubcf4 + \uc2dc\uac04 + \uc8fc\uc694 \ud544\ub4dc \uc21c\uc73c\ub85c '
    u'\uc6b0\uc120\uc21c\uc704\uac00 \ub192\uc740 \ud544\ub4dc\ubd80\ud130 \uc120\ud0dd\ud558\ub294 \uc2dc\ub098\ub9ac\uc624\ub97c \ubc18\uc601\ud55c\ub2e4.'
)

# Step2: 40컬럼
doc.add_heading(u'B.1 Step 2 - 40\uac1c \ucee8\ub7fc', level=2)
doc.add_paragraph(
    u'\uc804\uccb4 56\uac1c \uc911 \uc55e\uc5d0\uc11c 40\uac1c\ub97c \uc120\ud0dd\ud55c\ub2e4. '
    u'\uc7a5\ube44/\uc218\uc9d1 \uae30\ubcf8, \ub0a0\uc9dc/\uc2dc\uac04, \uc774\ubca4\ud2b8 \uc2dc\uac04, \ucd9c\ubc1c\uc9c0, \ubaa9\uc801\uc9c0 \uc815\ubcf4\uae4c\uc9c0 \ud3ec\ud568\ub41c\ub2e4. '
    u'\uc81c\uc678\ub418\ub294 16\uac1c: #41(dstnAssetNm) ~ #56(test1).'
)

# 40컬럼 테이블 (4열 x 10행)
cols40 = 40
t40 = doc.add_table(rows=1 + 10, cols=4)
t40.style = 'Light Grid Accent 1'
t40.alignment = WD_TABLE_ALIGNMENT.CENTER
for i in range(4):
    cell = t40.rows[0].cells[i]
    cell.text = [u'#1~10', u'#11~20', u'#21~30', u'#31~40'][i]
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(8)
for row_idx in range(10):
    for col_idx in range(4):
        field_idx = col_idx * 10 + row_idx
        cell = t40.rows[row_idx + 1].cells[col_idx]
        cell.text = f'{field_idx+1}. {all_fields[field_idx][1]}'
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
add_caption(doc, u'\ud45c B-1. Step 2 \uc120\ud0dd \ud544\ub4dc (40\uac1c)')

doc.add_paragraph()

# Step3: 20컬럼
doc.add_heading(u'B.2 Step 3 - 20\uac1c \ucee8\ub7fc', level=2)
doc.add_paragraph(
    u'\uc804\uccb4 56\uac1c \uc911 \uc55e\uc5d0\uc11c 20\uac1c\ub97c \uc120\ud0dd\ud55c\ub2e4. '
    u'\uc7a5\ube44/\uc218\uc9d1 \uae30\ubcf8 \uc815\ubcf4(11\uac1c) + \ub0a0\uc9dc/\uc2dc\uac04 \ubd84\ud574(8\uac1c) + \uc2dc\uc791\uc2dc\uac04(1\uac1c)\uc774 \ud3ec\ud568\ub41c\ub2e4. '
    u'\uc81c\uc678\ub418\ub294 36\uac1c: #21(endTime) ~ #56(test1).'
)

t20 = doc.add_table(rows=1 + 10, cols=2)
t20.style = 'Light Grid Accent 1'
t20.alignment = WD_TABLE_ALIGNMENT.CENTER
for i in range(2):
    cell = t20.rows[0].cells[i]
    cell.text = [u'#1~10', u'#11~20'][i]
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(8)
for row_idx in range(10):
    for col_idx in range(2):
        field_idx = col_idx * 10 + row_idx
        cell = t20.rows[row_idx + 1].cells[col_idx]
        cell.text = f'{field_idx+1}. {all_fields[field_idx][1]}'
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
add_caption(doc, u'\ud45c B-2. Step 3 \uc120\ud0dd \ud544\ub4dc (20\uac1c)')

doc.add_paragraph()

# Step4: 10컬럼
doc.add_heading(u'B.3 Step 4 - 10\uac1c \ucee8\ub7fc', level=2)
doc.add_paragraph(
    u'\uc804\uccb4 56\uac1c \uc911 \uc55e\uc5d0\uc11c 10\uac1c\ub9cc \uc120\ud0dd\ud55c\ub2e4. '
    u'\uc7a5\ube44/\uc218\uc9d1 \uae30\ubcf8 \uc815\ubcf4\uc758 \ud575\uc2ec \ud544\ub4dc\ub9cc \ud3ec\ud568\ub418\uba70, '
    u'\ub0a0\uc9dc \ubd84\ud574, \ub124\ud2b8\uc6cc\ud06c, \uc774\ubca4\ud2b8 \uc815\ubcf4\ub294 \ubaa8\ub450 \uc81c\uc678\ub41c\ub2e4.'
)

t10 = doc.add_table(rows=1 + 10, cols=2)
t10.style = 'Light Grid Accent 1'
t10.alignment = WD_TABLE_ALIGNMENT.CENTER
t10.rows[0].cells[0].text = u'#'
t10.rows[0].cells[1].text = u'\ud544\ub4dc\uba85'
for i in range(2):
    for p in t10.rows[0].cells[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
for row_idx in range(10):
    t10.rows[row_idx + 1].cells[0].text = str(row_idx + 1)
    t10.rows[row_idx + 1].cells[1].text = all_fields[row_idx][1]
    for c_idx in range(2):
        for p in t10.rows[row_idx + 1].cells[c_idx].paragraphs:
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
add_caption(doc, u'\ud45c B-3. Step 4 \uc120\ud0dd \ud544\ub4dc (10\uac1c)')

doc.add_paragraph()

# 프루닝 단계별 요약
doc.add_heading(u'B.4 \ub2e8\uacc4\ubcc4 \ud3ec\ud568 \uadf8\ub8f9 \uc694\uc57d', level=2)
summary_table = doc.add_table(rows=1 + 4, cols=5)
summary_table.style = 'Light Grid Accent 1'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sum_headers = [u'\ub2e8\uacc4', u'\ucee8\ub7fc \uc218', u'\ud3ec\ud568 \uadf8\ub8f9', u'\uc81c\uc678 \uc2dc\uc791', u'\uc5ed\uc9c1\ub82c\ud654 \uc2dc\uac04']
for i, h in enumerate(sum_headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(8)

sum_rows = [
    [u'Step 1', '56', u'\uc804\uccb4 (8\uac1c \uadf8\ub8f9)', '-', '27,823 ms'],
    [u'Step 2', '40', u'\uc7a5\ube44~\ubaa9\uc801\uc9c0 (5\uac1c \uadf8\ub8f9)', u'#41 dstnAssetNm', '19,453 ms'],
    [u'Step 3', '20', u'\uc7a5\ube44~\uc2dc\uc791\uc2dc\uac04 (3\uac1c \uadf8\ub8f9)', u'#21 endTime', '9,482 ms'],
    [u'Step 4', '10', u'\uc7a5\ube44/\uc218\uc9d1 \uae30\ubcf8 (1\uac1c \uadf8\ub8f9)', u'#11 parserKey', '4,609 ms'],
]
for r_idx, row in enumerate(sum_rows):
    for c_idx, val in enumerate(row):
        cell = summary_table.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        for p in cell.paragraphs:
            if c_idx in [1, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(8)

add_caption(doc, u'\ud45c B-4. \ucee8\ub7fc \ud504\ub8e8\ub2dd \ub2e8\uacc4\ubcc4 \ud3ec\ud568 \ubc94\uc704 \uc694\uc57d')

doc.add_page_break()

# ============================================================
# 부록 C. 쓰기 벤치마크 실측 데이터
# ============================================================
doc.add_heading(u'\ubd80\ub85d C. \uc4f0\uae30 \ubca4\uce58\ub9c8\ud06c \uc2e4\uce21 \ub370\uc774\ud130', level=1)
doc.add_paragraph(
    u'\uc4f0\uae30 \ubca4\uce58\ub9c8\ud06c\uc758 \ud3ec\ub9f7\ubcc4 \uac1c\ubcc4 \uc2e4\ud589 \uacb0\uacfc\uc774\ub2e4. '
    u'\ubaa8\ub4e0 \uc2e4\ud589\uc740 \ub3d9\uc77c\ud55c \ub370\uc774\ud130\uc14b(4,741,381\uac74, 5,427 MB)\uc744 \ub300\uc0c1\uc73c\ub85c \ud558\uc600\ub2e4.'
)

def add_small_table(doc, headers, rows, caption):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(7)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                if c_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(7)
    add_caption(doc, caption)
    return table

# 쓰기 실측 데이터 (에이전트에서 수집한 원본)
write_raw = {
    'tar': [
        [1, 46545, 104332, 9803, 183, 181594, 306762497],
        [2, 47299, 97671, 10248, 176, 177038, 306762497],
        [3, 49112, 97597, 10530, 189, 178026, 306762498],
        [4, 47004, 100195, 10699, 186, 178042, 306762498],
        [5, 50004, 103257, 10737, 189, 183606, 306762497],
        [6, 50293, 103696, 11073, 185, 187668, 306762497],
        [7, 47200, 96423, 10244, 184, 177690, 306762497],
        [8, 49374, 100651, 10382, 181, 179243, 306762498],
        [9, 47079, 96771, 11205, 169, 190518, 306762497],
        [10, 48750, 104193, 10740, 182, 182475, 306762498],
        [11, 51481, 97903, 11947, 188, 182410, 306762497],
    ],
    'kryo': [
        [1, 46090, 13181, 7266, 159, 79350, 278787788],
        [2, 46205, 12397, 8576, 165, 77841, 278787788],
        [3, 46378, 12940, 8091, 162, 78867, 278787788],
        [4, 47025, 13672, 8294, 159, 79862, 278787788],
        [5, 45676, 12097, 7258, 161, 76105, 278787788],
        [6, 46852, 12518, 7351, 163, 76831, 278787788],
        [7, 46520, 13164, 8275, 171, 78097, 278787788],
        [8, 46842, 13293, 7198, 156, 76847, 278787788],
        [9, 46606, 12405, 7216, 162, 75076, 278787788],
        [10, 45904, 13397, 7274, 157, 75841, 278787788],
        [11, 46812, 12369, 8295, 168, 80119, 278787788],
    ],
    'parquet': [
        [1, 47319, 71255, 5162, 120, 123953, 216465236],
        [2, 46886, 71470, 5672, 124, 124255, 216465236],
        [3, 47180, 71968, 6596, 120, 125964, 216465236],
        [4, 47087, 71664, 6753, 126, 125729, 216465236],
        [5, 48032, 72240, 4938, 128, 125440, 216465236],
        [6, 45817, 71595, 6331, 123, 123965, 216465236],
        [7, 46608, 71852, 6023, 126, 124709, 216465236],
        [8, 46332, 71157, 6953, 126, 124665, 216465238],
        [9, 45788, 71450, 6983, 124, 124445, 216465236],
        [10, 47252, 70049, 6612, 126, 124160, 216465208],
        [11, 45926, 70897, 5592, 129, 122707, 216465236],
    ],
    'avro': [
        [1, 46644, 29277, 6846, 144, 85075, 252357873],
        [2, 46744, 28198, 6882, 142, 85665, 252376271],
        [3, 46936, 28472, 6773, 144, 83711, 252373660],
        [4, 46814, 28763, 7056, 144, 85003, 252357836],
        [5, 46544, 29496, 8064, 146, 87463, 252377120],
        [6, 47192, 27998, 7034, 144, 84232, 252375558],
        [7, 47072, 29287, 7631, 149, 90329, 252356589],
        [8, 45114, 28870, 7498, 156, 86581, 252387272],
        [9, 46269, 29309, 7575, 157, 85952, 252364419],
        [10, 46603, 29547, 7410, 157, 92532, 252387267],
        [11, 46715, 29516, 7285, 153, 86010, 252354858],
    ],
}

wr_headers = [u'\ud68c\ucc28', u'\uc77d\uae30+\ud30c\uc2f1', u'\uc9c1\ub82c\ud654', u'\uc555\ucd95', u'\uc800\uc7a5', u'\uc804\uccb4(ms)', u'\uc555\ucd95\ud6c4(bytes)']
wr_ser_sizes = {'tar': '7,114,035,200', 'kryo': '2,841,639,885', 'parquet': '1,728,613,619', 'avro': '2,972,934,910'}

for fmt in FORMATS:
    doc.add_heading(f'C.{FORMATS.index(fmt)+1} {FORMAT_LABELS[fmt]}', level=2)
    doc.add_paragraph(
        u'\uc9c1\ub82c\ud654 \ud6c4 \uc6a9\ub7c9: ' + wr_ser_sizes[fmt] + u' bytes'
    )
    rows = []
    for run_data in write_raw[fmt]:
        rows.append([
            str(run_data[0]),
            f'{run_data[1]:,}',
            f'{run_data[2]:,}',
            f'{run_data[3]:,}',
            f'{run_data[4]:,}',
            f'{run_data[5]:,}',
            f'{run_data[6]:,}',
        ])
    # 평균 행
    n = len(write_raw[fmt])
    avgs = [sum(r[i] for r in write_raw[fmt]) / n for i in range(1, 6)]
    rows.append([
        u'\ud3c9\uade0',
        f'{avgs[0]:,.0f}',
        f'{avgs[1]:,.0f}',
        f'{avgs[2]:,.0f}',
        f'{avgs[3]:,.0f}',
        f'{avgs[4]:,.0f}',
        '-',
    ])
    t = add_small_table(doc, wr_headers, rows,
        u'\ud45c C-%d. %s \uc4f0\uae30 \ubca4\uce58\ub9c8\ud06c %d\ud68c \uc2e4\uce21 \uacb0\uacfc' % (FORMATS.index(fmt)+1, FORMAT_LABELS[fmt], n))
    # 평균 행 볼드
    last_row = t.rows[len(rows)]
    for cell in last_row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 부록 D. 읽기 벤치마크 실측 데이터
# ============================================================
doc.add_heading(u'\ubd80\ub85d D. \uc77d\uae30 \ubca4\uce58\ub9c8\ud06c \uc2e4\uce21 \ub370\uc774\ud130', level=1)
doc.add_paragraph(
    u'\uc77d\uae30 \ubca4\uce58\ub9c8\ud06c\uc758 \ud3ec\ub9f7/\ub2e8\uacc4\ubcc4 \uac1c\ubcc4 \uc2e4\ud589 \uacb0\uacfc\uc774\ub2e4. '
    u'\ubaa8\ub4e0 \uc2e4\ud589\uc740 \ub3d9\uc77c\ud55c .zst \ud30c\uc77c\uc744 \ud574\uc81c\ud55c \ud6c4 \uc5ed\uc9c1\ub82c\ud654\ud558\uc600\ub2e4. '
    u'\ubcf5\uc6d0 \ub85c\uadf8 \uac74\uc218\ub294 \ubaa8\ub450 4,741,381\uac74\uc774\ub2e4.'
)

read_raw = {
    'tar': {
        'step1': [
            [1, 98915, 45522, 145287],
            [2, 95299, 44554, 140615],
            [3, 101607, 43132, 145599],
            [4, 72480, 45114, 118436],
            [5, 78484, 44407, 123656],
        ],
        'step2': [
            [1, 72019, 45063, 117839],
            [2, 77186, 43430, 121368],
            [3, 72783, 43367, 116996],
            [4, 72503, 43719, 116963],
            [5, 89949, 44684, 135530],
        ],
    },
    'kryo': {
        'step1': [
            [1, 37355, 6732, 44429],
            [2, 36962, 6838, 44145],
            [3, 35511, 7029, 42845],
            [4, 38579, 6805, 45728],
            [5, 44215, 6867, 51428],
        ],
        'step2': [
            [1, 39662, 6763, 46737],
            [2, 43922, 6753, 50989],
            [3, 40490, 6865, 47717],
            [4, 49330, 6659, 56322],
            [5, 44230, 6877, 51465],
        ],
    },
    'parquet': {
        'step1': [
            [1, 21558, 27818, 49640],
            [2, 21336, 27778, 49291],
            [3, 19271, 27963, 47405],
            [4, 20172, 27897, 48252],
            [5, 18872, 27640, 46685],
        ],
        'step2': [
            [1, 20409, 19145, 39751],
            [2, 17987, 19428, 37588],
            [3, 21328, 19310, 40834],
            [4, 18969, 19539, 38689],
            [5, 18686, 19543, 38414],
        ],
        'step3': [
            [1, 18404, 9485, 28070],
            [2, 20883, 9500, 30577],
            [3, 23178, 9528, 32902],
            [4, 24173, 9378, 33768],
            [5, 24433, 9520, 34145],
        ],
        'step4': [
            [1, 27431, 4593, 32253],
            [2, 28658, 4607, 33457],
            [3, 29414, 4592, 34222],
            [4, 26272, 4617, 31113],
            [5, 27998, 4635, 32856],
        ],
    },
    'avro': {
        'step1': [
            [1, 42373, 23046, 65785],
            [2, 40036, 23468, 63828],
            [3, 38628, 23514, 62463],
            [4, 37111, 23656, 61127],
            [5, 40001, 24205, 64567],
        ],
        'step2': [
            [1, 40414, 16349, 57123],
            [2, 57748, 19357, 77396],
        ],
    },
}

rd_headers = [u'\ud68c\ucc28', u'\ud574\uc81c(zstd)', u'\uc5ed\uc9c1\ub82c\ud654', u'\uc804\uccb4(ms)']
step_names_base = {
    'step1': u'Step 1 - \uc804\uccb4 \uc77d\uae30',
    'step2': u'Step 2 - \ucd5c\uc801\ud654 \uc77d\uae30',
    'step3': u'Step 3 - 20\ucee8\ub7fc \ud504\ub8e8\ub2dd',
    'step4': u'Step 4 - 10\ucee8\ub7fc \ud504\ub8e8\ub2dd',
}
step_opt = {
    'tar':     {'step1': '', 'step2': u' (\ucd5c\uc801\ud654 \uc5c6\uc74c)'},
    'kryo':    {'step1': '', 'step2': u' (POJO \uc7ac\uc0ac\uc6a9)'},
    'parquet': {'step1': '', 'step2': u' (40\ucee8\ub7fc \ud504\ub8e8\ub2dd)', 'step3': '', 'step4': ''},
    'avro':    {'step1': '', 'step2': u' (SpecificDatumReader)'},
}
rd_comp_sizes = {'tar': '306,762,497', 'kryo': '278,787,788', 'parquet': '216,465,236', 'avro': '252,354,858'}
rd_decomp_sizes = {'tar': '7,114,035,200', 'kryo': '2,841,639,885', 'parquet': '1,728,613,619', 'avro': '2,972,934,910'}

table_counter = 0
for fmt in FORMATS:
    doc.add_heading(f'D.{FORMATS.index(fmt)+1} {FORMAT_LABELS[fmt]}', level=2)
    doc.add_paragraph(
        u'\uc555\ucd95 \ud30c\uc77c: ' + rd_comp_sizes[fmt] + u' bytes / '
        u'\ud574\uc81c \ud6c4: ' + rd_decomp_sizes[fmt] + u' bytes'
    )
    for step_key in sorted(read_raw[fmt].keys()):
        table_counter += 1
        step_data = read_raw[fmt][step_key]
        rows = []
        for run_data in step_data:
            rows.append([
                str(run_data[0]),
                f'{run_data[1]:,}',
                f'{run_data[2]:,}',
                f'{run_data[3]:,}',
            ])
        n = len(step_data)
        avgs = [sum(r[i] for r in step_data) / n for i in range(1, 4)]
        rows.append([
            u'\ud3c9\uade0',
            f'{avgs[0]:,.0f}',
            f'{avgs[1]:,.0f}',
            f'{avgs[2]:,.0f}',
        ])
        sname = step_names_base[step_key] + step_opt.get(fmt, {}).get(step_key, '')
        t = add_small_table(doc, rd_headers, rows,
            u'\ud45c D-%d. %s %s (%d\ud68c)' % (table_counter, FORMAT_LABELS[fmt], sname, n))
        last_row = t.rows[len(rows)]
        for cell in last_row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        doc.add_paragraph()

# ============================================================
# 저장
# ============================================================
output_path = '/opt/vortexlake_WR_bench3/document/VortexLake_Serialization_Benchmark_Report.docx'
doc.save(output_path)
print(f'Report saved: {output_path}')

import shutil
shutil.rmtree(tmpdir, ignore_errors=True)
